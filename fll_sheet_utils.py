import io
import pdfrw
from reportlab.pdfgen import canvas
import pandas as pd
import sys
import datetime
import docx
import qrcode
import icalendar
import uuid
import os


def roundTime(dt=None, dateDelta=datetime.timedelta(minutes=1)):
    """Round a datetime object to a multiple of a timedelta
    dt : datetime.datetime object, default now.
    dateDelta : timedelta object, we round to a multiple of this,
                default 1 minute.
    Author: Thierry Husson 2012, Stijn Nevens 2014
    """
    roundTo = dateDelta.total_seconds()

    if dt is None:
        dt = datetime.datetime.now()
    seconds = (dt - dt.min).seconds
    # // is a floor division, not a comment on following line:
    rounding = (seconds+roundTo/2) // roundTo * roundTo
    return dt + datetime.timedelta(0, rounding-seconds, -dt.microsecond)


def highlight_pit_table(table, pit_number):
    for row in range(len(table.rows)):
        for col in range(len(table.rows[0].cells)):
            if table.rows[row].cells[col].text == str(pit_number):
                # font = table.rows[row].cells[col].paragraphs[0].runs[0].font
                # font.bold = True
                # font.italic = False
                cell = table.rows[row].cells[col]
                cell.paragraphs[0].runs[0].style = 'Strong'


def rg_get_overlay_canvas(rg_round, team, table) -> io.BytesIO:
    data = io.BytesIO()
    pdf = canvas.Canvas(data)
    pdf.drawString(x=100, y=550, text=str(rg_round))
    pdf.drawString(x=100, y=575, text=str(team))
    pdf.drawString(x=290, y=550, text=table)
    pdf.save()
    data.seek(0)
    # print("Wrote overlay: round {0}, {1}, {2}".format(rg_round, team, table))
    return data


def rubric_get_overlay_canvas(team, place, teamloc, placeloc) -> io.BytesIO:
    data = io.BytesIO()
    pdf = canvas.Canvas(data)
    pdf.drawString(x=teamloc[0], y=teamloc[1], text=str(team))
    pdf.drawString(x=placeloc[0], y=placeloc[1], text=place)
    pdf.save()
    data.seek(0)
    # print("Wrote overlay: round {0}, {1}, {2}".format(rg_round, team, table))
    return data


def merge(overlay_canvas: io.BytesIO, template_path: str, page) -> io.BytesIO:
    template_pdf = pdfrw.PdfReader(template_path)
    overlay_pdf = pdfrw.PdfReader(overlay_canvas)
    data = overlay_pdf.pages[0]
    for page in [template_pdf.pages[page], ]:
        overlay = pdfrw.PageMerge().add(data)[0]
        pdfrw.PageMerge(page).add(overlay).render()
    form = io.BytesIO()
    pdfrw.PdfWriter().write(form, template_pdf)
    form.seek(0)
    return form


def save(form: io.BytesIO, filename: str):
    with open(filename, 'wb') as f:
        f.write(form.read())


def which_round(row, table):
    #    if row['Table1'] == table:
    #        return 1
    #    elif row['Table2'] == table:
    #        return 2
    #    elif row['Table3'] == table:
    #        return 3
    #    else:
    #        return None
    if pd.notna(row['Time1']):
        return 1
    elif pd.notna(row['Time2']):
        return 2
    elif pd.notna(row['Time3']):
        return 3
    else:
        return None


def which_time(row):
    if row['Round'] == 1:
        return row['Time1']
    elif row['Round'] == 2:
        return row['Time2']
    elif row['Round'] == 3:
        return row['Time3']
    else:
        return None


def make_qr_calendar(summary, start, duration, location):
    cal = icalendar.Calendar()
    qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
    )
    event = icalendar.Event()
    event.add('summary', summary)
    event.add('dtstart', start)
    event.add('dtend', start+datetime.timedelta(minutes=duration))
    event.add('location', location)
    cal.add_component(event)
    qr.add_data(cal.to_ical())
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    tf = '{0}.png'.format(str(uuid.uuid4()))
    img.save(tf)
    return tf


def make_referee_sheets(df, template):
    # Robot Game sheets
    rg_tables = pd.concat([df['Table1'],
                           df['Table2'],
                           df['Table3']]).unique()
    rg_tables.sort()
    print('Robot Game: ', end='', flush=True)
    for table in rg_tables:
        rg_writer = pdfrw.PdfWriter()
        rg_template_pdf = pdfrw.PdfReader(template)
        pdf = 'Robot Game {0}.pdf'.format(table)
        print('{0} '.format(table), end='', flush=True)
        # For each 8 tables and 48 teams, this probably comes down to:
        # 3x Round 1, 3x Round 2, 3x Round 1, 3x Round 2, 6x Round 3, but this
        # method should be generic for any 3-round tournament setup
#        teams_on_table = df[(df['Table1'] == table) |
#                            (df['Table2'] == table) |
#                            (df['Table3'] == table)]
        table_round_1 = df[df['Table1'] == table].sort_values('Time1')
        table_round_2 = df[df['Table2'] == table].sort_values('Time2')
        table_round_3 = df[df['Table3'] == table].sort_values('Time3')
        teams_on_table = pd.concat([table_round_1[['Team', 'Time1']],
                                    table_round_2[['Team', 'Time2']],
                                    table_round_3[['Team', 'Time3']]],
                                    sort=True)
        teams_on_table['Round'] = teams_on_table.apply(
                lambda row: which_round(row, table), axis=1)
        teams_on_table['Time'] = teams_on_table.apply(
                lambda row: which_time(row), axis=1)
        # Make a bunch of empty sheets for the merge
        print("({0} teams), ".format(len(teams_on_table)), end='', flush=True)
        for i in range(len(teams_on_table)):
            rg_writer.addpage(rg_template_pdf.pages[0])
        rg_writer.write(pdf)

        # This is a bit heavy on the reading/writing, but at 144 total pages,
        # it's not too bad so far.
        page_number = 0
        for index, row in teams_on_table.sort_values('Time').iterrows():
            print("{0} ".format(row['Team'].astype(int)), end='', flush=True)
            # base_pdf = pdfrw.PdfReader(pdf)
            canvas_data = rg_get_overlay_canvas(row['Round'].astype(int),
                                                row['Team'].astype(int),
                                                table)
            form = merge(canvas_data, template_path=pdf, page=page_number)
            page_number = page_number + 1
            save(form, filename=pdf)
        print(". ")
    print('')


def make_judge_sheets(df, template_path):
    rubric_template_pdf = pdfrw.PdfReader(template_path)

    # Core Values (page 1 from rubric template)
    cv_places = df['PlaceCV'].unique()
    cv_places.sort()
    template_page_number = 1
    print('Core Values: ', end='', flush=True)
    for place in cv_places:
        cv_writer = pdfrw.PdfWriter()
        pdf = 'Core Values {0}.pdf'.format(place)
        print('{0} '.format(place), end='', flush=True)
        teams_at_place = df[df['PlaceCV'] == place]
        for i in range(len(teams_at_place)):
            cv_writer.addpage(rubric_template_pdf.pages[
                    template_page_number-1])
        cv_writer.write(pdf)

        page_number = 0
        for index, row in teams_at_place.sort_values('TimeCV').iterrows():
            # base_pdf = pdfrw.PdfReader(pdf)
            canvas_data = rubric_get_overlay_canvas(row['Team'], place,
                                                    (480, 750), (480, 730))
            form = merge(canvas_data, template_path=pdf, page=page_number)
            page_number = page_number + 1
            save(form, filename=pdf)
    print('')

    # Innovation Project (page 2 from rubric template)
    ip_places = df['PlaceIP'].unique()
    ip_places.sort()
    template_page_number = 2
    print('Innovation Project: ', end='', flush=True)
    for place in ip_places:
        ip_writer = pdfrw.PdfWriter()
        pdf = 'Innovation Project {0}.pdf'.format(place)
        print('{0} '.format(place), end='', flush=True)
        teams_at_place = df[df['PlaceIP'] == place]
        for i in range(len(teams_at_place)):
            ip_writer.addpage(rubric_template_pdf.pages[
                    template_page_number-1])
        ip_writer.write(pdf)

        page_number = 0
        for index, row in teams_at_place.sort_values('TimeIP').iterrows():
            # base_pdf = pdfrw.PdfReader(pdf)
            canvas_data = rubric_get_overlay_canvas(row['Team'], place,
                                                    (480, 757), (480, 737))
            form = merge(canvas_data, template_path=pdf, page=page_number)
            page_number = page_number + 1
            save(form, filename=pdf)
    print('')

    # Robot Design (page 3 from rubric template)
    rd_places = df['PlaceRD'].unique()
    rd_places.sort()
    template_page_number = 3
    print('Robot Design: ', end='', flush=True)
    for place in rd_places:
        rd_writer = pdfrw.PdfWriter()
        pdf = 'Robot Design {0}.pdf'.format(place)
        print('{0} '.format(place), end='', flush=True)
        teams_at_place = df[df['PlaceRD'] == place]
        for i in range(len(teams_at_place)):
            rd_writer.addpage(rubric_template_pdf.pages[
                    template_page_number-1])
        rd_writer.write(pdf)

        page_number = 0
        for index, row in teams_at_place.sort_values('TimeIP').iterrows():
            # base_pdf = pdfrw.PdfReader(pdf)
            canvas_data = rubric_get_overlay_canvas(row['Team'], place,
                                                    (480, 750), (480, 730))
            form = merge(canvas_data, template_path=pdf, page=page_number)
            page_number = page_number + 1
            save(form, filename=pdf)
    print('')


def make_team_sheets(df, template_path, timeformat, buildingip, buildingrd,
                     buildingcv, date):
    today = datetime.datetime.strptime('{0} '.format(date)+'00:00:00',
                                       '%Y-%m-%d %H:%M:%S')
    pit_table_index = 0
    robot_game_table_index = 1
    judging_table_index = 2
    practice_round_count = 0
    for k in df.keys():
        if 'Practice' in k:
            # Practice rounds available, increment robot game and judging
            # table indices by 1
            practice_round_count += 1
            practice_table_index = 1
            robot_game_table_index = 2
            judging_table_index = 3

    for row in df.itertuples():

        pit = row.Pit
        name = row._3
        number = row.Team
        timep_dt = roundTime(dt=today+datetime.timedelta(row.TimePractice),
                             dateDelta=datetime.timedelta(minutes=5))
        timep = datetime.datetime.strftime(timep_dt, timeformat)
        tablep = row.TablePractice
        time1_dt = roundTime(dt=today+datetime.timedelta(row.Time1),
                             dateDelta=datetime.timedelta(minutes=5))
        time1 = datetime.datetime.strftime(time1_dt, timeformat)
        table1 = row.Table1
        time2_dt = roundTime(dt=today+datetime.timedelta(row.Time2),
                             dateDelta=datetime.timedelta(minutes=5))
        time2 = datetime.datetime.strftime(time2_dt, timeformat)
        table2 = row.Table2
        time3_dt = roundTime(dt=today+datetime.timedelta(row.Time3),
                             dateDelta=datetime.timedelta(minutes=5))
        time3 = datetime.datetime.strftime(time3_dt, timeformat)
        table3 = row.Table3
        timeip_dt = roundTime(dt=today+datetime.timedelta(row.TimeIP),
                              dateDelta=datetime.timedelta(minutes=5))
        timeip = datetime.datetime.strftime(timeip_dt, timeformat)
        placeip = row.PlaceIP
        timerd_dt = roundTime(dt=today+datetime.timedelta(row.TimeRD),
                              dateDelta=datetime.timedelta(minutes=5))
        timerd = datetime.datetime.strftime(timerd_dt, timeformat)
        placerd = row.PlaceRD
        timecv_dt = roundTime(dt=today+datetime.timedelta(row.TimeCV),
                              dateDelta=datetime.timedelta(minutes=5))
        timecv = datetime.datetime.strftime(timecv_dt, timeformat)
        placecv = row.PlaceCV

        document = docx.Document(template_path)
        rg_table = document.tables[robot_game_table_index]
        judging_table = document.tables[judging_table_index]
        document.paragraphs[0].text = 'Welcome, Team {0} ({1})!'.format(number,
                                                                        name)
        document.paragraphs[1].text = ('Your Pit Assignment: '
                                       'Pit {0}').format(pit)
        highlight_pit_table(document.tables[pit_table_index], pit)

        if practice_round_count > 0:
            # Practice Round
            prac_table = document.tables[practice_table_index]
            prac_desc_loc = prac_table.rows[1].cells[0]
            prac_desc_loc.text = 'Round 1: {0} on table {1}'.format(timep,
                                                                    tablep)
            qrp = make_qr_calendar(summary="Practice Round 1",
                                   start=timep_dt,
                                   duration=5,
                                   location="Table "+tablep)
            prac_qr_loc = prac_table.rows[0].cells[0].paragraphs[0]
            prac_qr_loc.add_run()
            prac_qr_loc.runs[0].add_picture(qrp,
                                            width=docx.shared.Inches(1.25))
            os.remove(qrp)

        # Robot Game Round 1
        rg1_desc_loc = rg_table.rows[1].cells[0]
        rg1_desc_loc.text = 'Round 1: {0} on table {1}'.format(time1, table1)
        qr1 = make_qr_calendar(summary="Robot Game Round 1",
                               start=time1_dt,
                               duration=5,
                               location="Table "+table1)
        rg1_qr_loc = rg_table.rows[0].cells[0].paragraphs[0]
        rg1_qr_loc.add_run()
        rg1_qr_loc.runs[0].add_picture(qr1,
                                       width=docx.shared.Inches(1.25))
        os.remove(qr1)

        # Robot Game Round 2
        rg2_desc_loc = rg_table.rows[1].cells[1]
        rg2_desc_loc.text = 'Round 2: {0} on table {1}'.format(time2, table2)
        qr2 = make_qr_calendar(summary="Robot Game Round 2",
                               start=time2_dt,
                               duration=5,
                               location="Table "+table2)
        rg2_qr_loc = rg_table.rows[0].cells[1].paragraphs[0]
        rg2_qr_loc.add_run()
        rg2_qr_loc.runs[0].add_picture(qr2,
                                       width=docx.shared.Inches(1.25))
        os.remove(qr2)

        # Robot Game Round 3
        rg3_desc_loc = rg_table.rows[1].cells[2]
        rg3_desc_loc.text = 'Round 3: {0} on table {1}'.format(time3, table3)
        qr3 = make_qr_calendar(summary="Robot Game Round 3",
                               start=time3_dt,
                               duration=5,
                               location="Table "+table3)
        rg3_qr_loc = rg_table.rows[0].cells[2].paragraphs[0]
        rg3_qr_loc.add_run()
        rg3_qr_loc.runs[0].add_picture(qr3,
                                       width=docx.shared.Inches(1.25))
        os.remove(qr3)

        # Innovation Project
        ip_desc_loc = judging_table.rows[1].cells[0]
        if buildingip is None:
            location = '{0}'.format(placeip)
            ip_desc_loc.text = ('Innovation Project: '
                                '{0} in {1}').format(timeip, placeip)
        else:
            location = '{0} {1}'.format(buildingip, placeip)
            ip_desc_loc.text = ('Innovation Project: '
                                '{0} in {1}, {2}').format(timeip, buildingip,
                                                          placeip)
        qr4 = make_qr_calendar(summary="Innovation Project Judging",
                               start=timeip_dt,
                               duration=10,
                               location=location)
        ip_qr_loc = judging_table.rows[0].cells[0].paragraphs[0]
        ip_qr_loc.add_run()
        ip_qr_loc.runs[0].add_picture(qr4,
                                      width=docx.shared.Inches(1.25))
        os.remove(qr4)

        # Robot Design
        rd_desc_loc = judging_table.rows[1].cells[1]
        if buildingrd is None:
            location = '{0}'.format(placerd)
            rd_desc_loc.text = ('Robot Design: '
                                '{0} in {1}').format(timerd, placerd)
        else:
            location = '{0} {1}'.format(buildingrd, placerd)
            rd_desc_loc.text = ('Robot Design: '
                                '{0} in {1}, {2}').format(timerd, buildingrd,
                                                          placerd)
        qr5 = make_qr_calendar(summary="Robot Design Judging",
                               start=timerd_dt,
                               duration=10,
                               location=location)
        rd_qr_loc = judging_table.rows[0].cells[1].paragraphs[0]
        rd_qr_loc.add_run()
        rd_qr_loc.runs[0].add_picture(qr5,
                                      width=docx.shared.Inches(1.25))
        os.remove(qr5)

        # Core Values
        cv_desc_loc = judging_table.rows[1].cells[2]
        if buildingcv is None:
            location = '{0}'.format(placecv)
            cv_desc_loc.text = ('Core Values: '
                                '{0} in {1}').format(timecv, placecv)
        else:
            location = '{0} {1}'.format(buildingcv, placecv)
            cv_desc_loc.text = ('Core Values: '
                                '{0} in {1}, {2}').format(timecv, buildingcv,
                                                          placecv)
        qr6 = make_qr_calendar(summary="Core Values Judging",
                               start=timecv_dt,
                               duration=10,
                               location=location)
        cv_qr_loc = judging_table.rows[0].cells[2].paragraphs[0]
        cv_qr_loc.add_run()
        cv_qr_loc.runs[0].add_picture(qr6,
                                      width=docx.shared.Inches(1.25))
        os.remove(qr6)

        print("Saving Team {0} ({1}).docx".format(number, name))
        document.save("Team {0} ({1}).docx".format(number, name))


def parse_args():
    try:
        in_file = open(sys.argv[1], 'r')
        in_file.close()
    except IndexError:
        sys.exit("Usage: {0} schedule.csv ".format(sys.argv[0]) +
                 "robot-game-template.pdf rubric.pdf " +
                 "team-template.docx YYYY-MM-DD")
    except IOError:
        sys.exit("Can't read CSV schedule {0}".format(sys.argv[1]))
    try:
        in_file = open(sys.argv[2], 'r')
        in_file.close()
    except IndexError:
        sys.exit("Usage: {0} schedule.csv ".format(sys.argv[0]) +
                 "robot-game-template.pdf rubric.pdf " +
                 "team-template.docx YYYY-MM-DD")
    except IOError:
        sys.exit("Can't read robot game template {0}".format(sys.argv[2]))
    try:
        in_file = open(sys.argv[3], 'r')
        in_file.close()
    except IndexError:
        sys.exit("Usage: {0} schedule.csv ".format(sys.argv[0]) +
                 "robot-game-template.pdf rubric.pdf " +
                 "team-template.docx YYYY-MM-DD")
    except IOError:
        sys.exit("Can't read rubric template {0}".format(sys.argv[3]))
    try:
        date = datetime.datetime.strptime(sys.argv[4], '%Y-%m-%d')
    except IndexError:
        sys.exit("Usage: {0} schedule.csv ".format(sys.argv[0]) +
                 "robot-game-template.pdf rubric.pdf " +
                 "team-template.docx YYYY-MM-DD")
    except ValueError:
        sys.exit("{0} is not in YYYY-MM-DD format".format(sys.argv[4]))
    schedule = sys.argv[1]
    rg_template_path = sys.argv[2]
    rubric_template_path = sys.argv[3]
    team_template_path = sys.argv[4]
    return (schedule, rg_template_path,
            rubric_template_path, team_template_path, date)


def read_csv(csv, column_types):
    pd.options.mode.chained_assignment = None
    return pd.read_table(csv, sep=',',
                         header=0,
                         names=column_types.keys(),
                         dtype=column_types)

