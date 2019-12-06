#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Nov 10 14:57:01 2019

@author: renfro
"""

import pandas as pd
import sys
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement


def roundTime(dt=None, dateDelta=datetime.timedelta(minutes=1)):
    """Round a datetime object to a multiple of a timedelta
    dt : datetime.datetime object, default now.
    dateDelta : timedelta object, we round to a multiple of this, default 1 minute.
    Author: Thierry Husson 2012, Stijn Nevens 2014
    """
    roundTo = dateDelta.total_seconds()

    if dt is None:
        dt = datetime.datetime.now()
    seconds = (dt - dt.min).seconds
    # // is a floor division, not a comment on following line:
    rounding = (seconds+roundTo/2) // roundTo * roundTo
    return dt + datetime.timedelta(0, rounding-seconds, -dt.microsecond)


def highlight_pit_table(table, pit_number):
    for row in range(len(table.rows)):
        for col in range(len(document.tables[0].rows[0].cells)):
            if table.rows[row].cells[col].text == str(pit_number):
                font = table.rows[row].cells[col].paragraphs[0].runs[0].font
                font.bold = True
                font.italic = True


column_types = {'Pit': str,
                'Team': int,
                'Team Name': str,
                'Time1': float,
                'Table1': str,
                'Time2': float,
                'Table2': str,
                'Time3': float,
                'Table3': str,
                'TimeIP': float,
                'PlaceIP': str,
                'TimeRD': float,
                'PlaceRD': str,
                'TimeCV': float,
                'PlaceCV': str}
schedule = 'fll-total-schedule.csv'
data = pd.read_table(schedule, sep=',',
                     header=0,
                     names=column_types.keys(),
                     dtype=column_types)

today = datetime.datetime.strptime('00:00:00', '%H:%M:%S')
timeformat = '%-I:%M %p'
buildingip = 'Daniels Hall'
buildingrd = 'Gym Basement'
buildingcv = 'Henderson Hall'

for row in data.itertuples():
    pit = row.Pit
    name = row._3
    number = row.Team
    time1 = roundTime(dt=today+datetime.timedelta(row.Time1),
                      dateDelta=datetime.timedelta(minutes=5))
    time1 = datetime.datetime.strftime(time1, timeformat)
    table1 = row.Table1
    time2 = roundTime(dt=today+datetime.timedelta(row.Time2),
                      dateDelta=datetime.timedelta(minutes=5))
    time2 = datetime.datetime.strftime(time2, timeformat)
    table2 = row.Table2
    time3 = roundTime(dt=today+datetime.timedelta(row.Time3),
                      dateDelta=datetime.timedelta(minutes=5))
    time3 = datetime.datetime.strftime(time3, timeformat)
    table3 = row.Table3
    timeip = roundTime(dt=today+datetime.timedelta(row.TimeIP),
                       dateDelta=datetime.timedelta(minutes=5))
    timeip = datetime.datetime.strftime(timeip, timeformat)
    placeip = row.PlaceIP
    timerd = roundTime(dt=today+datetime.timedelta(row.TimeRD),
                       dateDelta=datetime.timedelta(minutes=5))
    timerd = datetime.datetime.strftime(timerd, timeformat)
    placerd = row.PlaceRD
    timecv = roundTime(dt=today+datetime.timedelta(row.TimeCV),
                       dateDelta=datetime.timedelta(minutes=5))
    timecv = datetime.datetime.strftime(timecv, timeformat)
    placecv = row.PlaceCV

    document = docx.Document('team-template.docx')
    document.paragraphs[0].text = 'Welcome, Team {0} ({1})!'.format(number,
                                                                   name)
    document.paragraphs[1].text = 'Your Pit Assignment: Pit {0}'.format(pit)
    highlight_pit_table(document.tables[0], pit)

    document.add_heading('Your Robot Game Schedule', level=1)
    document.add_paragraph(
            'Round 1 at {0} on table {1}'.format(time1, table1),
            style='List Bullet')
    document.add_paragraph(
            'Round 2 at {0} on table {1}'.format(time2, table2),
            style='List Bullet')
    document.add_paragraph(
            'Round 3 at {0} on table {1}'.format(time3, table3),
            style='List Bullet')
    document.add_heading('Your Judging Schedule', level=1)
    document.add_paragraph(
            'Innovation Project at {0} '
            'in {1}, {2}'.format(timeip, buildingip, placeip),
            style='List Bullet')
    document.add_paragraph(
            'Robot Design at {0} '
            'in {1}, {2}'.format(timerd, buildingrd, placerd),
            style='List Bullet')
    document.add_paragraph(
            'Core Values at {0} '
            'in {1}, {2}'.format(timecv, buildingcv, placecv),
            style='List Bullet')

    print("Saving Team {0} ({1}).docx".format(number, name))
    document.save("Team {0} ({1}).docx".format(number, name))
